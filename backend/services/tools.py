"""
LLM tool definitions + executors.

Gemini'ye 3 araç tanıtıyoruz:
  - create_xlsx → openpyxl ile çok-sayfalı Excel
  - create_docx → python-docx ile Word belgesi
  - create_pdf  → reportlab ile basit markdown'lı PDF rapor

Üretilen dosya `outputs/` altına yazılır; her dosyaya UUID atanır,
frontend `GET /files/{file_id}` endpoint'inden indirir.
"""
from __future__ import annotations

import io
import os
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from google.genai import types as genai_types
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# ─── Output dizini + dosya kayıt ──────────────────────────────────────────────

_PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = Path(os.environ.get("NEURO_OUTPUT_DIR", str(_PROJECT_ROOT / "outputs")))
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# In-memory dosya indeksi (uvicorn yeniden başlayınca sıfırlanır — geçici dosyalar)
_FILE_INDEX: dict[str, dict] = {}


def _safe_name(name: str, ext: str) -> str:
    n = (name or "rapor").strip().replace("/", "_").replace("\\", "_")
    if not n.lower().endswith("." + ext):
        n = f"{n}.{ext}"
    return n


def _store(content: bytes, filename: str, mime: str) -> dict:
    file_id = str(uuid.uuid4())
    path = OUTPUT_DIR / f"{file_id}__{filename}"
    path.write_bytes(content)
    record = {"file_id": file_id, "filename": filename, "mime": mime, "path": str(path)}
    _FILE_INDEX[file_id] = record
    return {"file_id": file_id, "filename": filename, "mime": mime}


def get_file_info(file_id: str) -> dict | None:
    return _FILE_INDEX.get(file_id)


def get_file_path(file_id: str) -> Path | None:
    info = _FILE_INDEX.get(file_id)
    return Path(info["path"]) if info else None


# ─── Executors ────────────────────────────────────────────────────────────────


def _create_xlsx(filename: str, sheets: list[dict]) -> dict:
    wb = Workbook()
    # Default boş sheet'i sil; tüm sayfaları biz açacağız
    wb.remove(wb.active)
    if not sheets:
        sheets = [{"name": "Sheet1", "headers": [], "rows": []}]
    for sheet in sheets:
        title = str(sheet.get("name") or "Sheet")[:31]
        ws = wb.create_sheet(title=title)
        headers = sheet.get("headers") or []
        if headers:
            ws.append([str(h) for h in headers])
        for row in sheet.get("rows") or []:
            ws.append([str(c) if c is not None else "" for c in row])
    out = io.BytesIO()
    wb.save(out)
    return _store(
        out.getvalue(),
        _safe_name(filename, "xlsx"),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def _create_docx(filename: str, title: str, sections: list[dict]) -> dict:
    doc = DocxDocument()
    if title:
        doc.add_heading(title, level=0)
    for s in sections or []:
        heading = (s.get("heading") or "").strip()
        body = (s.get("body") or "").strip()
        if heading:
            doc.add_heading(heading, level=1)
        if body:
            for paragraph in body.split("\n\n"):
                doc.add_paragraph(paragraph)
    out = io.BytesIO()
    doc.save(out)
    return _store(
        out.getvalue(),
        _safe_name(filename, "docx"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _create_pdf(filename: str, title: str, markdown: str) -> dict:
    """Basit markdown alt-kümesi: # / ## başlıklar, - / * listeler, satır = paragraf."""
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4, title=title or filename)
    styles = getSampleStyleSheet()
    story = []
    if title:
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 12))

    for line in (markdown or "").splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 8))
            continue
        if stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], styles["Heading1"]))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], styles["Heading2"]))
        elif stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], styles["Heading3"]))
        elif stripped.startswith(("- ", "* ")):
            story.append(Paragraph("• " + stripped[2:], styles["BodyText"]))
        else:
            story.append(Paragraph(stripped, styles["BodyText"]))
    if not story:
        story.append(Paragraph("(Boş içerik)", styles["BodyText"]))
    doc.build(story)
    return _store(out.getvalue(), _safe_name(filename, "pdf"), "application/pdf")


_EXECUTORS = {
    "create_xlsx": _create_xlsx,
    "create_docx": _create_docx,
    "create_pdf": _create_pdf,
}


def execute_tool(name: str, args: dict) -> dict:
    fn = _EXECUTORS.get(name)
    if not fn:
        return {"error": f"Bilinmeyen araç: {name}"}
    try:
        return fn(**args)
    except TypeError as exc:
        return {"error": f"Araç parametre hatası ({name}): {exc}"}
    except Exception as exc:  # pragma: no cover
        return {"error": f"Araç çalıştırılırken hata ({name}): {exc}"}


# ─── Gemini function declarations ─────────────────────────────────────────────


def _str_array() -> genai_types.Schema:
    return genai_types.Schema(
        type=genai_types.Type.ARRAY,
        items=genai_types.Schema(type=genai_types.Type.STRING),
    )


def get_tool_declarations() -> list[genai_types.Tool]:
    return [
        genai_types.Tool(
            function_declarations=[
                genai_types.FunctionDeclaration(
                    name="create_xlsx",
                    description=(
                        "Excel (.xlsx) dosyası oluşturur ve indirilebilir hale getirir. "
                        "Tablo verilerini, raporları, listeleri kullanıcıya sunmak için kullan. "
                        "Birden çok sayfa destekler."
                    ),
                    parameters=genai_types.Schema(
                        type=genai_types.Type.OBJECT,
                        properties={
                            "filename": genai_types.Schema(
                                type=genai_types.Type.STRING,
                                description="Dosya adı (uzantısız).",
                            ),
                            "sheets": genai_types.Schema(
                                type=genai_types.Type.ARRAY,
                                items=genai_types.Schema(
                                    type=genai_types.Type.OBJECT,
                                    properties={
                                        "name": genai_types.Schema(type=genai_types.Type.STRING),
                                        "headers": _str_array(),
                                        "rows": genai_types.Schema(
                                            type=genai_types.Type.ARRAY,
                                            items=_str_array(),
                                        ),
                                    },
                                    required=["name", "rows"],
                                ),
                            ),
                        },
                        required=["filename", "sheets"],
                    ),
                ),
                genai_types.FunctionDeclaration(
                    name="create_docx",
                    description=(
                        "Word (.docx) belgesi oluşturur. Yazılı raporlar, dokümanlar, "
                        "mektuplar ve uzun metinler için kullan."
                    ),
                    parameters=genai_types.Schema(
                        type=genai_types.Type.OBJECT,
                        properties={
                            "filename": genai_types.Schema(type=genai_types.Type.STRING),
                            "title": genai_types.Schema(type=genai_types.Type.STRING),
                            "sections": genai_types.Schema(
                                type=genai_types.Type.ARRAY,
                                items=genai_types.Schema(
                                    type=genai_types.Type.OBJECT,
                                    properties={
                                        "heading": genai_types.Schema(type=genai_types.Type.STRING),
                                        "body": genai_types.Schema(type=genai_types.Type.STRING),
                                    },
                                    required=["body"],
                                ),
                            ),
                        },
                        required=["filename", "title", "sections"],
                    ),
                ),
                genai_types.FunctionDeclaration(
                    name="create_pdf",
                    description=(
                        "PDF raporu oluşturur. Sunum, özet, formal raporlar için kullan. "
                        "İçerik markdown formatında verilir (#, ##, ###, -, *)."
                    ),
                    parameters=genai_types.Schema(
                        type=genai_types.Type.OBJECT,
                        properties={
                            "filename": genai_types.Schema(type=genai_types.Type.STRING),
                            "title": genai_types.Schema(type=genai_types.Type.STRING),
                            "markdown": genai_types.Schema(type=genai_types.Type.STRING),
                        },
                        required=["filename", "title", "markdown"],
                    ),
                ),
            ]
        )
    ]
